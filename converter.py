from __future__ import annotations

import re
import sys
import tempfile
import urllib.parse
import urllib.request
from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

import markdown
from docx import Document
from docx.document import Document as _DocumentType
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from latex2mathml.converter import convert as latex_to_mathml
from lxml import etree, html as lxml_html
from PIL import Image

TOKEN_PATTERN = re.compile(r"(MATH(?:BLOCK|INLINE)TOK\d+END)")
CAPTION_PREFIX_PATTERN = re.compile(
    r"^(?:(?:表|图)\s*\d+(?:\s*[:：.．]\s*|\s+)\S|(?:table|figure)\s*\d+(?:\s*[:.]\s*|\s+)\S)",
    re.IGNORECASE,
)
INLINE_MATH_TOKEN = "MATHINLINE"
BLOCK_MATH_TOKEN = "MATHBLOCK"
MATHML_NS = "http://www.w3.org/1998/Math/MathML"
MATHML_TAG = f"{{{MATHML_NS}}}"
SCRIPT_MATH_CHARS = {
    "𝒜": "A",
    "ℬ": "B",
    "𝒞": "C",
    "𝒟": "D",
    "ℰ": "E",
    "ℱ": "F",
    "𝒢": "G",
    "ℋ": "H",
    "ℐ": "I",
    "𝒥": "J",
    "𝒦": "K",
    "ℒ": "L",
    "ℳ": "M",
    "𝒩": "N",
    "𝒪": "O",
    "𝒫": "P",
    "𝒬": "Q",
    "ℛ": "R",
    "𝒮": "S",
    "𝒯": "T",
    "𝒰": "U",
    "𝒱": "V",
    "𝒲": "W",
    "𝒳": "X",
    "𝒴": "Y",
    "𝒵": "Z",
}

LATIN_FONT = "Times New Roman"
CJK_FONT = "宋体"
CODE_FONT = "Consolas"
ARTICLE_TITLE_STYLE = "文章标题"
TOC_TITLE_STYLE = "TOC 标题"

ARTICLE_TITLE_SIZE_PT = 16.0
TOC_TITLE_SIZE_PT = 16.0
H1_SIZE_PT = 14.0
H2_SIZE_PT = 12.0
H3_SIZE_PT = 12.0
H4_SIZE_PT = 12.0
H5_SIZE_PT = 12.0
H6_SIZE_PT = 12.0
BODY_SIZE_PT = 12.0
CAPTION_SIZE_PT = 10.5
CODE_SIZE_PT = 10.5
BODY_FIRST_LINE_INDENT_PT = 24.0
LIST_INDENT_PT = 18.0
LIST_MARKER_MIN_WIDTH_PT = 11.0
LIST_MARKER_CHAR_WIDTH_PT = 5.5
BLOCKQUOTE_LEFT_INDENT_PT = 0.0
BLOCKQUOTE_RIGHT_INDENT_PT = 3.0
BLOCKQUOTE_FIRST_LINE_INDENT_PT = 0.0
BLOCKQUOTE_SPACE_BEFORE_PT = 3.0
BLOCKQUOTE_SPACE_AFTER_PT = 3.0
BLOCKQUOTE_BORDER_SPACE_PT = 3.0
BLOCKQUOTE_BORDER_SIZE = "12"
BLOCKQUOTE_BORDER_COLOR = "A6A6A6"
BLOCKQUOTE_SHADING = "F2F2F2"


@dataclass
class AppConfig:
    output_dir: str
    asset_root: str
    title_chars: int = 12
    auto_timestamp: bool = True
    body_first_line_indent: bool = True


@dataclass
class TextStyle:
    bold: bool = False
    italic: bool = False
    underline: bool = False
    code: bool = False
    size_pt: float = BODY_SIZE_PT

    def copy(self) -> "TextStyle":
        return TextStyle(
            bold=self.bold,
            italic=self.italic,
            underline=self.underline,
            code=self.code,
            size_pt=self.size_pt,
        )


class MarkdownToDocxConverter:
    def __init__(self, xsl_path: Optional[str] = None) -> None:
        root = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))
        self.xsl_path = Path(xsl_path) if xsl_path else root / "mml2omml.xsl"
        if not self.xsl_path.exists():
            raise FileNotFoundError(f"Cannot find XSLT file: {self.xsl_path}")

        self._math_transform = etree.XSLT(etree.parse(str(self.xsl_path)))
        self._math_tokens: Dict[str, Dict[str, object]] = {}
        self.last_warnings: List[str] = []
        self._document: Optional[_DocumentType] = None

    def convert(self, markdown_text: str, config: AppConfig) -> str:
        if not markdown_text or not markdown_text.strip():
            raise ValueError("Markdown text is empty.")

        output_dir = Path(config.output_dir).expanduser().resolve()
        output_dir.mkdir(parents=True, exist_ok=True)

        self.last_warnings = []
        prepared_markdown = self._prepare_markdown(markdown_text)
        html_text = markdown.markdown(
            prepared_markdown,
            extensions=[
                "tables",
                "fenced_code",
                "sane_lists",
                "md_in_html",
            ],
            output_format="html5",
        )

        document = Document()
        self._document = document
        self._configure_document_styles(document)

        root = lxml_html.fragment_fromstring(html_text, create_parent="div")
        for child in root:
            self._render_block(child, document, config)

        output_path = output_dir / self._build_output_filename(markdown_text, config)
        document.save(str(output_path))
        return str(output_path)

    def latex_to_omml(self, latex: str, display: bool) -> etree._Element:
        normalized_latex = self._preprocess_latex_for_math(latex, display)
        mathml_text = latex_to_mathml(normalized_latex)
        mathml_root = etree.fromstring(mathml_text.encode("utf-8"))
        mathml_root.set("display", "block" if display else "inline")
        self._normalize_mathml_tree(mathml_root)

        transformed = self._math_transform(mathml_root)
        if transformed.getroot() is None:
            raise ValueError("Math transform returned empty result.")
        omml_root = deepcopy(transformed.getroot())
        self._normalize_omml_math(omml_root, display=display)
        if re.search(r"\\begin\{(?:aligned|align\*?|split)\}", latex):
            self._apply_aligned_matrix_layout(omml_root)
        if re.search(r"\\begin\{cases\}", latex):
            self._apply_cases_matrix_layout(omml_root)
        return omml_root

    def _preprocess_latex_for_math(self, latex: str, display: bool) -> str:
        normalized = latex

        normalized = re.sub(r"\\\\\s*\[[^\]]+\]", r"\\\\", normalized)
        normalized = self._normalize_evaluation_bar_after_fraction(normalized)
        normalized = self._convert_aligned_environments_to_arrays(normalized)
        if re.search(r"\\begin\{(?:aligned|align\*?|split)\}", normalized):
            normalized = re.sub(r"(?<!\\)&", "", normalized)

        # `\pmb` should preserve the original italic math style for variables.
        normalized = re.sub(r"\\pmb(?=\s*\{)", r"\\boldsymbol", normalized)

        if display:
            large_ops = r"sum|prod|coprod|bigcup|bigcap|bigoplus|bigotimes|bigvee|bigwedge|bigsqcup"
            normalized = re.sub(
                rf"\\({large_ops})(?!\s*\\(?:limits|nolimits))(?=\s*(?:_|\^))",
                lambda match: f"\\{match.group(1)}\\limits",
                normalized,
            )

        return normalized

    def _normalize_evaluation_bar_after_fraction(self, latex: str) -> str:
        pattern = re.compile(
            r"(\\frac\{(?:[^{}]|\{[^{}]*\})+\}\{(?:[^{}]|\{[^{}]*\})+\})"
            r"\s*\\(?:Big|big|Bigg|bigg)?\|(_\{[^{}]+\})"
        )
        return pattern.sub(r"\\left.\1\\right|\2", latex)

    def _convert_aligned_environments_to_arrays(self, latex: str) -> str:
        pattern = re.compile(r"\\begin\{(?P<env>aligned|align\*?|split)\}(?P<body>.*?)\\end\{(?P=env)\}", re.DOTALL)

        def replacer(match: re.Match[str]) -> str:
            body = match.group("body")
            rows = re.split(r"(?<!\\)\\\\", body)
            max_columns = max((len(re.findall(r"(?<!\\)&", row)) + 1 for row in rows), default=2)
            max_columns = max(2, max_columns)
            column_spec = "r" + ("l" * (max_columns - 1))
            return f"\\begin{{array}}{{{column_spec}}}{body}\\end{{array}}"

        return pattern.sub(replacer, latex)

    def _normalize_mathml_tree(self, root: etree._Element) -> None:
        self._normalize_mathml_script_chars(root)
        self._convert_mathml_spaces(root)
        self._convert_mathml_stretchy_fences(root)

    def _normalize_mathml_script_chars(self, root: etree._Element) -> None:
        for node in root.xpath(".//*[local-name()='mi']"):
            text = (node.text or "").strip()
            if text in SCRIPT_MATH_CHARS:
                node.text = SCRIPT_MATH_CHARS[text]
                node.set("mathvariant", "script")

    def _convert_mathml_spaces(self, root: etree._Element) -> None:
        for node in root.xpath(".//*[local-name()='mspace']"):
            width = (node.get("width") or "").strip().lower()
            replacement = etree.Element(f"{MATHML_TAG}mtext")
            replacement.text = self._mathml_space_text(width)

            parent = node.getparent()
            if parent is not None:
                parent.replace(node, replacement)

    def _convert_mathml_stretchy_fences(self, node: etree._Element) -> None:
        for child in list(node):
            self._convert_mathml_stretchy_fences(child)

        if self._mathml_local_name(node) != "mrow":
            return

        children = list(node)
        if len(children) < 2:
            return

        first = children[0]
        last = children[-1]
        if not self._is_mathml_stretchy_fence(first, "prefix"):
            return
        if not self._is_mathml_stretchy_fence(last, "postfix"):
            return

        mfenced = etree.Element(f"{MATHML_TAG}mfenced")
        mfenced.set("open", self._mathml_text(first))
        mfenced.set("close", self._mathml_text(last))
        mfenced.set("separators", "")

        inner_row = etree.Element(f"{MATHML_TAG}mrow")
        for child in children[1:-1]:
            inner_row.append(child)
        mfenced.append(inner_row)

        parent = node.getparent()
        if parent is not None:
            parent.replace(node, mfenced)

    def _is_mathml_stretchy_fence(self, node: etree._Element, expected_form: str) -> bool:
        return (
            self._mathml_local_name(node) == "mo"
            and (node.get("stretchy") or "").lower() == "true"
            and (node.get("fence") or "").lower() == "true"
            and (node.get("form") or "").lower() == expected_form
        )

    def _mathml_local_name(self, node: etree._Element) -> str:
        return etree.QName(node).localname.lower()

    def _mathml_text(self, node: etree._Element) -> str:
        return "".join(node.itertext())

    def _mathml_space_text(self, width: str) -> str:
        if width.endswith("em"):
            try:
                amount = float(width[:-2])
            except ValueError:
                amount = 1.0
        else:
            amount = 1.0

        if amount >= 1.5:
            return "\u2003\u2003"
        if amount >= 0.75:
            return "\u2003"
        if amount >= 0.4:
            return "\u2002"
        return " "

    def _normalize_omml_math(self, omml_root: etree._Element, display: bool) -> None:
        for math_node in self._iter_omath_nodes(omml_root):
            self._repair_nary_operand(math_node)
            if not display:
                self._convert_inline_nary_to_scripts(math_node)
            else:
                self._convert_display_nary_runs(math_node)
            self._repair_matrix_delimiter(math_node)
            self._repair_leading_matrix_delimiter(math_node)
            self._tune_nary_style(math_node)

    def _tune_nary_style(self, omath: etree._Element) -> None:
        for nary in omath.xpath(".//*[local-name()='nary']"):
            grow_nodes = nary.xpath("./*[local-name()='naryPr']/*[local-name()='grow']")
            for grow in grow_nodes:
                grow.set(qn("m:val"), "0")

    def _convert_inline_nary_to_scripts(self, omath: etree._Element) -> None:
        nary_symbols = {"∑", "∏", "∐", "⋃", "⋂", "⨁", "⨂", "∨", "∧", "⨆"}

        for nary in list(omath.xpath(".//*[local-name()='nary']")):
            symbol = self._nary_symbol(nary)
            if symbol not in nary_symbols:
                continue

            parent = nary.getparent()
            if parent is None:
                continue

            sub_node = self._first_child_by_local_name(nary, "sub")
            sup_node = self._first_child_by_local_name(nary, "sup")
            e_node = self._first_child_by_local_name(nary, "e")
            if sub_node is None and sup_node is None:
                continue

            script_node = self._build_inline_nary_script(symbol, sub_node, sup_node)
            insert_index = parent.index(nary)
            parent.replace(nary, script_node)

            if e_node is None:
                continue
            for operand_child in list(e_node):
                insert_index += 1
                parent.insert(insert_index, operand_child)

    def _convert_display_nary_runs(self, omath: etree._Element) -> None:
        nary_symbols = {"∑", "∏", "∐", "⋃", "⋂", "⨁", "⨂", "∨", "∧", "⨆"}
        changed = True

        while changed:
            changed = False
            for parent in omath.iter():
                children = list(parent)
                for idx, child in enumerate(children):
                    if self._tag_name(child) != "r":
                        continue

                    symbol = self._get_run_text(child).strip()
                    if symbol not in nary_symbols:
                        continue

                    operand_nodes = self._collect_nary_operand_siblings(parent, idx + 1)
                    if not operand_nodes:
                        continue

                    nary = self._build_display_nary(symbol, operand_nodes)
                    parent.replace(child, nary)
                    changed = True
                    break
                if changed:
                    break

    def _build_display_nary(self, symbol: str, operand_nodes: List[etree._Element]) -> etree._Element:
        nary = OxmlElement("m:nary")
        nary_pr = OxmlElement("m:naryPr")

        chr_node = OxmlElement("m:chr")
        chr_node.set(qn("m:val"), symbol)
        lim_loc = OxmlElement("m:limLoc")
        lim_loc.set(qn("m:val"), "undOvr")
        grow = OxmlElement("m:grow")
        grow.set(qn("m:val"), "0")
        sub_hide = OxmlElement("m:subHide")
        sub_hide.set(qn("m:val"), "on")
        sup_hide = OxmlElement("m:supHide")
        sup_hide.set(qn("m:val"), "on")

        nary_pr.append(chr_node)
        nary_pr.append(lim_loc)
        nary_pr.append(grow)
        nary_pr.append(sub_hide)
        nary_pr.append(sup_hide)
        nary.append(nary_pr)

        e_node = OxmlElement("m:e")
        for operand_node in operand_nodes:
            e_node.append(operand_node)
        nary.append(e_node)
        return nary

    def _nary_symbol(self, nary: etree._Element) -> str:
        chr_nodes = nary.xpath("./*[local-name()='naryPr']/*[local-name()='chr']")
        if not chr_nodes:
            return ""
        return chr_nodes[0].get(qn("m:val"), "")

    def _first_child_by_local_name(self, node: etree._Element, local_name: str) -> Optional[etree._Element]:
        for child in node:
            if self._tag_name(child) == local_name:
                return child
        return None

    def _build_inline_nary_script(
        self,
        symbol: str,
        sub_node: Optional[etree._Element],
        sup_node: Optional[etree._Element],
    ) -> etree._Element:
        if sub_node is not None and sup_node is not None:
            script = OxmlElement("m:sSubSup")
        elif sub_node is not None:
            script = OxmlElement("m:sSub")
        else:
            script = OxmlElement("m:sSup")

        base = OxmlElement("m:e")
        run = OxmlElement("m:r")
        text = OxmlElement("m:t")
        text.text = symbol
        run.append(text)
        base.append(run)
        script.append(base)

        if sub_node is not None:
            script.append(deepcopy(sub_node))
        if sup_node is not None:
            script.append(deepcopy(sup_node))

        return script

    def _iter_omath_nodes(self, root: etree._Element):
        local = self._tag_name(root)
        if local == "omath":
            yield root
            return
        if local == "omathpara":
            for node in root.xpath("./*[local-name()='oMath']"):
                yield node

    def _repair_nary_operand(self, omath: etree._Element) -> None:
        changed = True
        while changed:
            changed = False
            for parent in omath.iter():
                if self._repair_nary_operand_in_parent(parent):
                    changed = True
                    break

    def _repair_nary_operand_in_parent(self, parent: etree._Element) -> bool:
        children = list(parent)
        for idx, child in enumerate(children):
            if self._tag_name(child) != "nary":
                continue
            e_nodes = child.xpath("./*[local-name()='e']")
            if not e_nodes:
                continue
            e_node = e_nodes[0]
            if len(e_node) != 0:
                continue
            if idx + 1 >= len(children):
                continue

            collected = self._collect_nary_operand_siblings(parent, idx + 1)
            if not collected:
                continue
            for operand_node in collected:
                e_node.append(operand_node)
            return True
        return False

    def _collect_nary_operand_siblings(self, parent: etree._Element, start_index: int) -> List[etree._Element]:
        collected: List[etree._Element] = []

        while start_index < len(parent):
            sibling = parent[start_index]
            tag = self._tag_name(sibling)
            if tag in {"dPr", "ctrlPr", "nary"}:
                break

            if tag == "r":
                text = self._get_run_text(sibling)
                if self._starts_with_nary_stop_operator(text):
                    break

                prefix, remainder = self._split_trailing_nary_stop_operator(text)
                if prefix:
                    operand_run = deepcopy(sibling)
                    self._set_run_text(operand_run, prefix)
                    collected.append(operand_run)

                    if remainder:
                        self._set_run_text(sibling, remainder)
                        break
                    parent.remove(sibling)
                    continue

                break

            collected.append(sibling)
            parent.remove(sibling)

        return collected

    def _starts_with_nary_stop_operator(self, text: str) -> bool:
        return bool(text and text.lstrip().startswith(("+", "-", "−", "±", "∓")))

    def _split_trailing_nary_stop_operator(self, text: str) -> Tuple[str, str]:
        match = re.search(r"([+\-−±∓])\s*$", text or "")
        if not match:
            return text, ""
        prefix = text[: match.start(1)].rstrip()
        remainder = text[match.start(1) :]
        return prefix, remainder

    def _repair_matrix_delimiter(self, omath: etree._Element) -> None:
        bracket_pairs = {"(": ")", "[": "]", "{": "}", "|": "|"}
        target_tags = {"m", "eqarr"}
        changed = True

        while changed:
            changed = False
            children = list(omath)
            for idx, expr_node in enumerate(children):
                if self._tag_name(expr_node) not in target_tags:
                    continue
                if idx == 0 or idx >= len(children) - 1:
                    continue

                left_node = children[idx - 1]
                right_node = children[idx + 1]
                left = self._peek_open_bracket_from_run(left_node)
                right = self._peek_close_bracket_from_run(right_node)

                if left is None or right is None:
                    continue
                left_bracket, left_remain = left
                right_bracket, right_remain = right
                if bracket_pairs.get(left_bracket) != right_bracket:
                    continue

                delim = self._build_delimiter_omml(left_bracket, right_bracket, expr_node)
                omath.replace(expr_node, delim)

                self._set_run_text(left_node, left_remain)
                self._set_run_text(right_node, right_remain)
                if not left_remain:
                    omath.remove(left_node)
                if not right_remain:
                    omath.remove(right_node)

                changed = True
                break

    def _repair_leading_matrix_delimiter(self, omath: etree._Element) -> None:
        target_tags = {"m", "eqarr"}
        changed = True

        while changed:
            changed = False
            children = list(omath)
            for idx, expr_node in enumerate(children):
                if self._tag_name(expr_node) not in target_tags or idx == 0:
                    continue

                left_node = children[idx - 1]
                left = self._peek_open_bracket_from_run(left_node)
                if left is None:
                    continue

                left_bracket, left_remain = left
                if left_bracket not in {"{", "[", "(", "|"}:
                    continue

                delim = self._build_delimiter_omml(left_bracket, "", expr_node)
                omath.replace(expr_node, delim)

                self._set_run_text(left_node, left_remain)
                if not left_remain:
                    omath.remove(left_node)

                changed = True
                break

    def _apply_aligned_matrix_layout(self, omml_root: etree._Element) -> None:
        for matrix in omml_root.xpath(".//*[local-name()='m']"):
            column_count = self._matrix_column_count(matrix)
            if column_count <= 1:
                continue

            alignments = ["right"] + ["left"] * (column_count - 1)
            if column_count >= 3:
                alignments[-1] = "right"
            self._set_matrix_column_alignments(matrix, alignments)

    def _apply_cases_matrix_layout(self, omml_root: etree._Element) -> None:
        for matrix in omml_root.xpath(".//*[local-name()='m']"):
            column_count = self._matrix_column_count(matrix)
            if column_count <= 0:
                continue
            self._set_matrix_column_alignments(matrix, ["left"] * column_count)

    def _matrix_column_count(self, matrix: etree._Element) -> int:
        counts = [len(row.xpath("./*[local-name()='e']")) for row in matrix.xpath("./*[local-name()='mr']")]
        return max(counts, default=0)

    def _set_matrix_column_alignments(self, matrix: etree._Element, alignments: List[str]) -> None:
        matrix_pr = None
        for child in matrix:
            if self._tag_name(child) == "mpr":
                matrix_pr = child
                break

        if matrix_pr is None:
            matrix_pr = OxmlElement("m:mPr")
            matrix.insert(0, matrix_pr)

        for child in list(matrix_pr):
            if self._tag_name(child) == "mcs":
                matrix_pr.remove(child)

        mcs = OxmlElement("m:mcs")
        for alignment in alignments:
            mc = OxmlElement("m:mc")
            mc_pr = OxmlElement("m:mcPr")
            count = OxmlElement("m:count")
            count.set(qn("m:val"), "1")
            mc_jc = OxmlElement("m:mcJc")
            mc_jc.set(qn("m:val"), alignment)
            mc_pr.append(count)
            mc_pr.append(mc_jc)
            mc.append(mc_pr)
            mcs.append(mc)
        matrix_pr.append(mcs)

    def _peek_open_bracket_from_run(self, node: etree._Element) -> Optional[Tuple[str, str]]:
        if self._tag_name(node) != "r":
            return None
        text = self._get_run_text(node)
        if not text:
            return None

        stripped = text.rstrip()
        if not stripped:
            return None
        bracket = stripped[-1]
        if bracket not in {"(", "[", "{", "|"}:
            return None
        remain = stripped[:-1] + text[len(stripped) :]
        return bracket, remain

    def _peek_close_bracket_from_run(self, node: etree._Element) -> Optional[Tuple[str, str]]:
        if self._tag_name(node) != "r":
            return None
        text = self._get_run_text(node)
        if not text:
            return None

        leading_trimmed = text.lstrip()
        if not leading_trimmed:
            return None
        bracket = leading_trimmed[0]
        if bracket not in {")", "]", "}", "|"}:
            return None
        remain = text[: len(text) - len(leading_trimmed)] + leading_trimmed[1:]
        return bracket, remain

    def _get_run_text(self, node: etree._Element) -> str:
        return "".join((t.text or "") for t in node.xpath(".//*[local-name()='t']"))

    def _set_run_text(self, node: etree._Element, text: str) -> None:
        text_nodes = node.xpath("./*[local-name()='t']")
        if text_nodes:
            text_nodes[0].text = text
            for extra in text_nodes[1:]:
                node.remove(extra)
            return

        if text:
            t = OxmlElement("m:t")
            t.text = text
            node.append(t)

    def _is_operator_run(self, node: etree._Element) -> bool:
        text = "".join((t.text or "") for t in node.xpath(".//*[local-name()='t']")).strip()
        return text in {"+", "-", "*", "/", "=", ",", ";", ":"}

    def _build_delimiter_omml(self, left: str, right: str, expr_node: etree._Element) -> etree._Element:
        delim = OxmlElement("m:d")
        d_pr = OxmlElement("m:dPr")

        beg = OxmlElement("m:begChr")
        beg.set(qn("m:val"), left)
        d_pr.append(beg)

        end = OxmlElement("m:endChr")
        end.set(qn("m:val"), right)
        d_pr.append(end)

        grow = OxmlElement("m:grow")
        grow.set(qn("m:val"), "1")
        d_pr.append(grow)

        delim.append(d_pr)
        e = OxmlElement("m:e")
        e.append(deepcopy(expr_node))
        delim.append(e)
        return delim

    def resolve_image_path(self, src: str, asset_root: str) -> Optional[Path]:
        if not src:
            return None

        src = src.strip()
        if self._is_remote_url(src) or src.startswith("data:"):
            return None

        image_path = Path(src)
        if image_path.is_absolute() and image_path.exists():
            return image_path

        base_dir = Path(asset_root).expanduser().resolve() if asset_root else Path.cwd()
        candidate = (base_dir / image_path).resolve()
        if candidate.exists():
            return candidate

        filename_candidate = (base_dir / image_path.name).resolve() if image_path.name else None
        if filename_candidate is not None and filename_candidate.exists():
            return filename_candidate

        local_candidate = (Path.cwd() / image_path).resolve()
        if local_candidate.exists():
            return local_candidate

        return None

    def _configure_document_styles(self, document: _DocumentType) -> None:
        self._configure_style(document, "Normal", LATIN_FONT, CJK_FONT, BODY_SIZE_PT)
        self._configure_style(
            document,
            ARTICLE_TITLE_STYLE,
            LATIN_FONT,
            CJK_FONT,
            ARTICLE_TITLE_SIZE_PT,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            outline_level=9,
            bold=True,
            font_color=RGBColor(0x36, 0x5F, 0x91),
        )
        self._configure_style(
            document,
            TOC_TITLE_STYLE,
            LATIN_FONT,
            CJK_FONT,
            TOC_TITLE_SIZE_PT,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            outline_level=9,
            bold=True,
            font_color=RGBColor(0x36, 0x5F, 0x91),
        )
        self._configure_style(document, "Heading 1", LATIN_FONT, CJK_FONT, H1_SIZE_PT, outline_level=0)
        self._configure_style(document, "Heading 2", LATIN_FONT, CJK_FONT, H2_SIZE_PT, outline_level=1)
        self._configure_style(document, "Heading 3", LATIN_FONT, CJK_FONT, H3_SIZE_PT, outline_level=2)
        self._configure_style(
            document,
            "Heading 4",
            LATIN_FONT,
            CJK_FONT,
            H4_SIZE_PT,
            outline_level=3,
            font_color=RGBColor(0x24, 0x3F, 0x60),
        )
        self._configure_style(document, "Heading 5", LATIN_FONT, CJK_FONT, H5_SIZE_PT, outline_level=4)
        self._configure_style(document, "Heading 6", LATIN_FONT, CJK_FONT, H6_SIZE_PT, outline_level=5)
        self._configure_style(
            document,
            "Caption",
            LATIN_FONT,
            CJK_FONT,
            CAPTION_SIZE_PT,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent_pt=0,
            outline_level=9,
        )
        self._configure_hyperlink_style(document, "Hyperlink", RGBColor(5, 99, 193))
        self._configure_hyperlink_style(document, "FollowedHyperlink", RGBColor(149, 79, 114))

    def _configure_style(
        self,
        document: _DocumentType,
        style_name: str,
        latin_font: str,
        east_asia_font: str,
        size_pt: float,
        alignment: Optional[WD_ALIGN_PARAGRAPH] = None,
        first_line_indent_pt: Optional[float] = None,
        outline_level: Optional[int] = None,
        bold: Optional[bool] = None,
        font_color: Optional[RGBColor] = None,
    ) -> None:
        try:
            style = document.styles[style_name]
        except KeyError:
            style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

        style.font.name = latin_font
        style.font.size = Pt(size_pt)
        style.font.italic = False
        if bold is not None:
            style.font.bold = bold
        if font_color is not None:
            style.font.color.rgb = font_color
        if alignment is not None and style.type == WD_STYLE_TYPE.PARAGRAPH:
            style.paragraph_format.alignment = alignment
        if first_line_indent_pt is not None and style.type == WD_STYLE_TYPE.PARAGRAPH:
            style.paragraph_format.first_line_indent = Pt(first_line_indent_pt)

        try:
            style.hidden = False
            style.quick_style = True
            style.unhide_when_used = True
        except AttributeError:
            pass

        r_pr = style._element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:ascii"), latin_font)
        r_fonts.set(qn("w:hAnsi"), latin_font)
        r_fonts.set(qn("w:eastAsia"), east_asia_font)
        r_fonts.set(qn("w:cs"), latin_font)
        for theme_attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme", "w:csTheme"):
            attr_name = qn(theme_attr)
            if attr_name in r_fonts.attrib:
                del r_fonts.attrib[attr_name]
        if font_color is not None:
            color = r_pr.find(qn("w:color"))
            if color is None:
                color = OxmlElement("w:color")
                r_pr.append(color)
            color.set(qn("w:val"), str(font_color))
            for theme_attr in ("w:themeColor", "w:themeTint", "w:themeShade"):
                attr_name = qn(theme_attr)
                if attr_name in color.attrib:
                    del color.attrib[attr_name]

        if outline_level is not None and style.type == WD_STYLE_TYPE.PARAGRAPH:
            self._set_style_outline_level(style, outline_level)

    def _set_style_outline_level(self, style, outline_level: int) -> None:
        p_pr = style._element.get_or_add_pPr()
        existing = p_pr.find(qn("w:outlineLvl"))
        if existing is not None:
            p_pr.remove(existing)
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), str(outline_level))
        p_pr.append(outline)

    def _configure_hyperlink_style(self, document: _DocumentType, style_name: str, rgb_color: RGBColor) -> None:
        try:
            style = document.styles[style_name]
        except KeyError:
            style = document.styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)

        style.font.name = LATIN_FONT
        style.font.underline = True
        style.font.color.rgb = rgb_color

        r_pr = style._element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:ascii"), LATIN_FONT)
        r_fonts.set(qn("w:hAnsi"), LATIN_FONT)
        r_fonts.set(qn("w:eastAsia"), CJK_FONT)
        r_fonts.set(qn("w:cs"), LATIN_FONT)

    def _prepare_markdown(self, text: str) -> str:
        masked_text, code_tokens = self._mask_code_regions(text)
        indent_ready_text = self._normalize_two_space_list_indents(masked_text)
        boundary_ready_text = self._normalize_markdown_block_boundaries(indent_ready_text)
        list_ready_text = self._split_reset_ordered_lists(boundary_ready_text)
        caption_ready_text = self._merge_caption_headings(list_ready_text)
        math_ready_text = self._extract_math_tokens(caption_ready_text)
        return self._restore_tokens(math_ready_text, code_tokens)

    def _normalize_two_space_list_indents(self, text: str) -> str:
        normalized_lines: List[str] = []
        list_item_pattern = re.compile(r"^([ \t]*)([-+*]\s+|\d+[.)]\s+)(.*)$")

        for line in text.splitlines():
            match = list_item_pattern.match(line)
            if not match:
                normalized_lines.append(line)
                continue

            raw_indent, marker, rest = match.groups()
            indent_width = len(raw_indent.replace("\t", "    "))
            normalized_level = indent_width // 2
            normalized_lines.append(f"{' ' * (normalized_level * 4)}{marker}{rest}")

        return "\n".join(normalized_lines)

    def _normalize_markdown_block_boundaries(self, text: str) -> str:
        lines = text.splitlines()
        result: List[str] = []

        for line in lines:
            if result and self._needs_blank_before_markdown_block(line, result):
                result.append("")
            result.append(line)

        return "\n".join(result)

    def _needs_blank_before_markdown_block(self, line: str, previous_lines: List[str]) -> bool:
        if not line.strip():
            return False

        prev_index = len(previous_lines) - 1
        while prev_index >= 0 and not previous_lines[prev_index].strip():
            prev_index -= 1

        if prev_index < 0 or prev_index != len(previous_lines) - 1:
            return False

        previous = previous_lines[prev_index]
        if self._is_markdown_list_item(line):
            return not self._is_markdown_list_item(previous)

        if self._is_markdown_horizontal_rule(line):
            return True

        return False

    def _is_markdown_list_item(self, line: str) -> bool:
        return bool(re.match(r"^[ \t]*(?:[-+*]\s+|\d+[.)]\s+)", line))

    def _is_markdown_horizontal_rule(self, line: str) -> bool:
        return bool(re.match(r"^[ \t]{0,3}(-[ \t]*){3,}$", line))

    def _merge_caption_headings(self, text: str) -> str:
        lines = text.splitlines()
        result: List[str] = []
        index = 0

        while index < len(lines):
            match = re.match(r"^(#{2,6})\s+(.*?)\s*$", lines[index])
            if not match:
                result.append(lines[index])
                index += 1
                continue

            heading_text = re.sub(r"\s+#+\s*$", "", match.group(2)).strip()
            if not self._looks_like_caption_start(heading_text):
                result.append(lines[index])
                index += 1
                continue

            next_index = index + 1
            while next_index < len(lines) and not lines[next_index].strip():
                next_index += 1

            if next_index >= len(lines) or self._is_block_boundary(lines[next_index]):
                result.append(lines[index])
                index += 1
                continue

            paragraph_lines: List[str] = []
            end_index = next_index
            while end_index < len(lines) and lines[end_index].strip():
                if self._is_block_boundary(lines[end_index]) and paragraph_lines:
                    break
                paragraph_lines.append(lines[end_index])
                end_index += 1

            if not paragraph_lines:
                result.append(lines[index])
                index += 1
                continue

            merged_paragraph = self._collapse_markdown_lines(paragraph_lines)
            result.append(self._merge_caption_heading_and_body(heading_text, merged_paragraph))
            index = end_index

        return "\n".join(result)

    def _collapse_markdown_lines(self, lines: List[str]) -> str:
        combined = ""
        for line in lines:
            segment = line.strip()
            if not segment:
                continue
            if not combined:
                combined = segment
                continue

            if self._needs_join_space(combined[-1], segment[0]):
                combined += " "
            combined += segment
        return combined

    def _needs_join_space(self, left_char: str, right_char: str) -> bool:
        return bool(re.match(r"[A-Za-z0-9]", left_char) and re.match(r"[A-Za-z0-9]", right_char))

    def _merge_caption_heading_and_body(self, heading_text: str, body_text: str) -> str:
        heading_with_period = self._ensure_caption_terminal_punctuation(heading_text)
        separator = ""
        if re.search(r"[.!?]$", heading_with_period) and body_text and re.match(r"[A-Za-z0-9]", body_text[0]):
            separator = " "
        return f"{heading_with_period}{separator}{body_text}"

    def _ensure_caption_terminal_punctuation(self, text: str) -> str:
        if re.search(r"[。．.!！？?]$", text):
            return text
        if re.match(r"^(table|figure)\s*\d+(?:\s*[:.]\s*|\s+)\S", text.strip(), re.IGNORECASE):
            return f"{text}."
        return f"{text}。"

    def _is_block_boundary(self, line: str) -> bool:
        stripped = line.lstrip()
        return bool(
            re.match(r"^(#{1,6})\s+", stripped)
            or re.match(r"^(```|~~~)", stripped)
            or re.match(r"^>\s*", stripped)
            or re.match(r"^(\*|-|\+)\s+", stripped)
            or re.match(r"^\d+[.)]\s+", stripped)
            or stripped.startswith("|")
            or stripped.startswith("<")
        )

    def _looks_like_caption_start(self, text: str) -> bool:
        normalized = re.sub(r"\s+", " ", text or "").strip()
        return bool(CAPTION_PREFIX_PATTERN.match(normalized))

    def _mask_code_regions(self, text: str) -> Tuple[str, Dict[str, str]]:
        tokens: Dict[str, str] = {}

        def fenced_replacer(match: re.Match[str]) -> str:
            token = f"CODEBLOCKTOK{len(tokens)}END"
            tokens[token] = match.group(0)
            return token

        text = re.sub(r"```[\s\S]*?```|~~~[\s\S]*?~~~", fenced_replacer, text)

        def inline_replacer(match: re.Match[str]) -> str:
            token = f"CODEINLINETOK{len(tokens)}END"
            tokens[token] = match.group(0)
            return token

        text = re.sub(r"`[^`\n]+`", inline_replacer, text)
        return text, tokens

    def _restore_tokens(self, text: str, tokens: Dict[str, str]) -> str:
        restored = text
        for token, value in tokens.items():
            restored = restored.replace(token, value)
        return restored

    def _extract_math_tokens(self, text: str) -> str:
        self._math_tokens = {}

        def create_token(latex: str, display: bool) -> str:
            token_type = BLOCK_MATH_TOKEN if display else INLINE_MATH_TOKEN
            token = f"{token_type}TOK{len(self._math_tokens)}END"
            if display:
                latex = self._strip_blockquote_markers(latex)
            self._math_tokens[token] = {"latex": latex.strip(), "display": display}
            if display:
                return f"\n\n{token}\n\n"
            return token

        def blockquote_display_replacer(match: re.Match[str]) -> str:
            prefix = match.group("prefix")
            body = self._strip_blockquote_markers(match.group("body"))
            token = create_token(body, True).strip()
            return f"{prefix}\n{prefix}{token}\n{prefix}"

        blockquote_display_pattern = re.compile(
            r"(?ms)^(?P<prefix>[ \t]*>[ \t]*)\\\[[ \t]*\n"
            r"(?P<body>.*?)[ \t]*>[ \t]*\\\][ \t]*$"
        )
        text = blockquote_display_pattern.sub(blockquote_display_replacer, text)

        text = re.sub(r"\$\$(.+?)\$\$", lambda m: create_token(m.group(1), True), text, flags=re.DOTALL)
        text = re.sub(r"\\\[(.+?)\\\]", lambda m: create_token(m.group(1), True), text, flags=re.DOTALL)
        text = re.sub(r"\\\((.+?)\\\)", lambda m: create_token(m.group(1), False), text)
        text = re.sub(r"(?<!\\)\$(?!\$)([^\n]+?)(?<!\\)\$", lambda m: create_token(m.group(1), False), text)
        return text

    def _strip_blockquote_markers(self, text: str) -> str:
        lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
        return "\n".join(re.sub(r"^[ \t]*>[ \t]?", "", line) for line in lines).strip()

    def _split_reset_ordered_lists(self, text: str) -> str:
        ordered_item_pattern = re.compile(r"^(\s*)(\d+)([.)])\s+")
        separator = '<div class="md2docx-list-split"></div>'
        lines = text.splitlines()
        result: List[str] = []
        index = 0

        while index < len(lines):
            current = lines[index]
            if current.strip():
                result.append(current)
                index += 1
                continue

            blank_start = index
            while index < len(lines) and not lines[index].strip():
                result.append(lines[index])
                index += 1

            prev_index = blank_start - 1
            while prev_index >= 0 and not lines[prev_index].strip():
                prev_index -= 1

            if prev_index < 0 or index >= len(lines):
                continue

            prev_match = ordered_item_pattern.match(lines[prev_index])
            next_match = ordered_item_pattern.match(lines[index])
            if not prev_match or not next_match:
                continue

            same_indent = len(prev_match.group(1)) == len(next_match.group(1))
            restarts_at_one = next_match.group(2) == "1"
            if same_indent and restarts_at_one:
                result.append(separator)

        return "\n".join(result)

    def _render_block(self, node: etree._Element, container, config: AppConfig) -> None:
        tag = self._tag_name(node)
        if not tag:
            return

        if tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            level = int(tag[1])
            style_name = self._markdown_heading_style_name(level)
            para = container.add_paragraph(style=style_name)
            self._apply_alignment(para, node)
            if style_name == ARTICLE_TITLE_STYLE or self._is_abstract_title(node.text_content()):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = Pt(0)
            size = self._markdown_heading_size(level)
            self._render_inline_node(node, para, container, config, TextStyle(size_pt=size))
            return

        if tag in {"p", "div", "center"}:
            if not node.text_content().strip() and not node.xpath(".//*[local-name()='img']"):
                return

            single_block = self._single_block_math_token(node)
            if single_block:
                para = container.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = Pt(0)
                self._append_math(paragraph=para, token=single_block, inline=False)
                return

            para = container.add_paragraph()
            if tag == "center":
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = Pt(0)
            elif self._is_image_only_block(node) or self._is_abstract_title(node.text_content()):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = Pt(0)
            else:
                self._apply_alignment(para, node)

            style = TextStyle(size_pt=BODY_SIZE_PT)
            if self._is_caption_text(node.text_content()):
                try:
                    para.style = "Caption"
                except KeyError:
                    pass
                style.size_pt = CAPTION_SIZE_PT
                para.paragraph_format.first_line_indent = Pt(0)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif self._is_image_only_block(node) or self._is_abstract_title(node.text_content()):
                para.paragraph_format.first_line_indent = Pt(0)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif tag != "center" and config.body_first_line_indent:
                para.paragraph_format.first_line_indent = Pt(BODY_FIRST_LINE_INDENT_PT)
            else:
                para.paragraph_format.first_line_indent = Pt(0)

            self._render_inline_node(node, para, container, config, style)
            return

        if tag in {"ul", "ol"}:
            self._render_list(node, container, config, ordered=(tag == "ol"), level=1)
            return

        if tag == "pre":
            self._render_code_block(node, container)
            return

        if tag == "blockquote":
            start_index = len(getattr(container, "paragraphs", []))
            for child in node:
                self._render_block(child, container, config)
            for para in getattr(container, "paragraphs", [])[start_index:]:
                self._apply_blockquote_paragraph_format(para)
            return

        if tag == "table":
            self._render_table(node, container, config)
            return

        if tag == "img":
            para = container.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Pt(0)
            self._insert_image(para, node, config)
            return

        if tag == "hr":
            self._render_horizontal_rule(container)
            return

        for child in node:
            self._render_block(child, container, config)

    def _render_code_block(self, node: etree._Element, container) -> None:
        code_text = (node.text_content() or "").replace("\r\n", "\n").replace("\r", "\n")
        lines = code_text.split("\n")
        if lines and lines[-1] == "":
            lines = lines[:-1]
        if not lines:
            lines = [""]

        for idx, line in enumerate(lines):
            para = container.add_paragraph()
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.space_before = Pt(6 if idx == 0 else 0)
            para.paragraph_format.space_after = Pt(6 if idx == len(lines) - 1 else 0)
            self._set_paragraph_shading(para, "EDEDED")

            run = para.add_run(line if line else " ")
            self._apply_run_style(run, TextStyle(code=True, size_pt=CODE_SIZE_PT))

    def _render_list(self, node: etree._Element, container, config: AppConfig, ordered: bool, level: int) -> None:
        next_number = self._extract_list_start(node) if ordered else 1

        for li in node.xpath("./li"):
            para = container.add_paragraph()
            task_state = None if ordered else self._extract_task_state(li)
            marker = f"{next_number}." if ordered else self._task_marker(task_state) if task_state is not None else self._bullet_prefix(level)
            prefix = f"{marker} \t"
            self._configure_list_paragraph(para, level, marker_text=f"{marker} ")
            prefix_run = para.add_run(prefix)
            self._apply_run_style(prefix_run, TextStyle(size_pt=BODY_SIZE_PT))
            self._render_list_item_inline(li, para, container, config)
            if ordered:
                next_number += 1

            for nested in li:
                nested_tag = self._tag_name(nested)
                if nested_tag in {"ul", "ol"}:
                    self._render_list(nested, container, config, ordered=(nested_tag == "ol"), level=level + 1)

    def _render_list_item_inline(self, li: etree._Element, paragraph, container, config: AppConfig) -> None:
        base_style = TextStyle(size_pt=BODY_SIZE_PT)
        self._render_inline_node(li, paragraph, container, config, base_style)

    def _extract_task_state(self, li: etree._Element) -> Optional[bool]:
        checkbox = li.find("./input")
        if checkbox is not None and (checkbox.get("type") or "").lower() == "checkbox":
            checkbox.getparent().remove(checkbox)
            return checkbox.get("checked") is not None

        text_owner = li
        first_text = li.text or ""
        if not first_text.strip():
            first_paragraph = li.find("./p")
            if first_paragraph is not None:
                text_owner = first_paragraph
                first_text = first_paragraph.text or ""

        task_text = self._normalize_inline_text(first_text)
        match = re.match(r"^\s*\[([xX ])\]\s*", task_text)
        if not match:
            return None

        text_owner.text = re.sub(r"^\s*\[[xX ]\]\s*", "", first_text, count=1)
        return match.group(1).lower() == "x"

    def _task_marker(self, checked: Optional[bool]) -> str:
        return "☑" if checked else "☐"

    def _render_table(self, table_node: etree._Element, container, config: AppConfig) -> None:
        rows = table_node.xpath("./thead/tr|./tbody/tr|./tfoot/tr|./tr")
        if not rows:
            return

        occupancy: Dict[Tuple[int, int], bool] = {}
        anchors: Dict[Tuple[int, int], Tuple[etree._Element, int, int]] = {}
        max_col = 0
        max_row = 0

        for r_idx, row in enumerate(rows):
            c_idx = 0
            cells = row.xpath("./th|./td")
            for cell in cells:
                while occupancy.get((r_idx, c_idx), False):
                    c_idx += 1

                rowspan = max(1, int(cell.get("rowspan", "1") or "1"))
                colspan = max(1, int(cell.get("colspan", "1") or "1"))
                anchors[(r_idx, c_idx)] = (cell, rowspan, colspan)

                for rr in range(r_idx, r_idx + rowspan):
                    for cc in range(c_idx, c_idx + colspan):
                        occupancy[(rr, cc)] = True
                        max_row = max(max_row, rr)
                        max_col = max(max_col, cc)

                c_idx += colspan

        table = container.add_table(rows=max_row + 1, cols=max_col + 1)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for (r_idx, c_idx), (_, rowspan, colspan) in anchors.items():
            if rowspan > 1 or colspan > 1:
                table.cell(r_idx, c_idx).merge(table.cell(r_idx + rowspan - 1, c_idx + colspan - 1))

        for (r_idx, c_idx), (cell_node, _, _) in anchors.items():
            cell = table.cell(r_idx, c_idx)
            self._clear_cell(cell)
            para = cell.add_paragraph()
            style = TextStyle(bold=(self._tag_name(cell_node) == "th"), size_pt=BODY_SIZE_PT)
            self._render_inline_node(cell_node, para, cell, config, style)

    def _render_inline_node(self, node: etree._Element, paragraph, container, config: AppConfig, style: TextStyle):
        if node.text and (node.text.strip() or len(node) == 0):
            paragraph = self._append_text(paragraph, container, node.text, style)

        for child in node:
            child_tag = self._tag_name(child)
            child_style = style.copy()

            if child_tag in {"ul", "ol"}:
                continue
            if child_tag == "input" and (child.get("type") or "").lower() == "checkbox":
                continue
            if child_tag in {"strong", "b"}:
                child_style.bold = True
                paragraph = self._render_inline_node(child, paragraph, container, config, child_style)
            elif child_tag in {"em", "i"}:
                child_style.italic = True
                paragraph = self._render_inline_node(child, paragraph, container, config, child_style)
            elif child_tag == "u":
                child_style.underline = True
                paragraph = self._render_inline_node(child, paragraph, container, config, child_style)
            elif child_tag == "code":
                child_style.code = True
                child_style.size_pt = CODE_SIZE_PT
                paragraph = self._render_inline_node(child, paragraph, container, config, child_style)
            elif child_tag == "a":
                paragraph = self._append_hyperlink_from_node(paragraph, container, child, child_style, config)
            elif child_tag == "br":
                paragraph = self._start_new_paragraph_after(paragraph, container)
            elif child_tag == "img":
                self._insert_image(paragraph, child, config)
            elif child_tag in {"sub", "sup"}:
                paragraph = self._append_script_text(
                    paragraph,
                    container,
                    child.text_content(),
                    child_style,
                    is_subscript=(child_tag == "sub"),
                    is_superscript=(child_tag == "sup"),
                )
            else:
                paragraph = self._render_inline_node(child, paragraph, container, config, child_style)

            tail_text = child.tail or ""
            if child_tag == "br" and tail_text.startswith("\n"):
                tail_text = tail_text[1:]
            if tail_text and (tail_text.strip() or self._tag_name(node) in {"code", "pre"}):
                paragraph = self._append_text(paragraph, container, tail_text, style)

        return paragraph

    def _append_hyperlink_from_node(self, paragraph, container, anchor_node: etree._Element, style: TextStyle, config: AppConfig):
        href = (anchor_node.get("href") or "").strip()
        text = self._normalize_inline_text(anchor_node.text_content()).strip() or href

        if not href:
            return self._append_text(paragraph, container, text, style)
        resolved_href = self._resolve_hyperlink_target(href, config.asset_root)

        for part in self._split_text_with_paragraphs(text):
            if part == "\n":
                paragraph = self._start_new_paragraph_after(paragraph, container)
                continue
            if not part:
                continue
            if part in self._math_tokens:
                self._append_math(paragraph=paragraph, token=part, inline=not self._math_tokens[part]["display"])
                continue
            self._append_hyperlink(paragraph, resolved_href, part, style)
        return paragraph

    def _resolve_hyperlink_target(self, href: str, asset_root: str) -> str:
        href = href.strip()
        if not href:
            return href
        if href.startswith("#"):
            return href

        path_like = Path(href)
        if path_like.is_absolute():
            try:
                return path_like.resolve().as_uri()
            except ValueError:
                return str(path_like)

        if re.match(r"^[a-zA-Z][a-zA-Z0-9+.-]*:", href):
            return href

        base_dir = Path(asset_root).expanduser().resolve() if asset_root else Path.cwd()
        resolved = (base_dir / path_like).resolve()
        try:
            return resolved.as_uri()
        except ValueError:
            return str(resolved)

    def _append_hyperlink(self, paragraph, url: str, text: str, style: TextStyle) -> None:
        part = paragraph.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        hyperlink.set(qn("w:history"), "1")

        run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")

        r_style = OxmlElement("w:rStyle")
        r_style.set(qn("w:val"), "Hyperlink")
        r_pr.append(r_style)

        r_fonts = OxmlElement("w:rFonts")
        if style.code:
            r_fonts.set(qn("w:ascii"), CODE_FONT)
            r_fonts.set(qn("w:hAnsi"), CODE_FONT)
            r_fonts.set(qn("w:eastAsia"), CODE_FONT)
            r_fonts.set(qn("w:cs"), CODE_FONT)
        else:
            r_fonts.set(qn("w:ascii"), LATIN_FONT)
            r_fonts.set(qn("w:hAnsi"), LATIN_FONT)
            r_fonts.set(qn("w:eastAsia"), CJK_FONT)
            r_fonts.set(qn("w:cs"), LATIN_FONT)
        r_pr.append(r_fonts)

        size_half = str(int(round(style.size_pt * 2)))
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), size_half)
        r_pr.append(sz)
        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), size_half)
        r_pr.append(sz_cs)

        if style.bold:
            r_pr.append(OxmlElement("w:b"))
        if style.italic:
            r_pr.append(OxmlElement("w:i"))
        if style.underline:
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            r_pr.append(u)
        if style.code:
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "EDEDED")
            r_pr.append(shd)

        run.append(r_pr)
        text_element = OxmlElement("w:t")
        text_element.text = text
        run.append(text_element)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    def _append_text(self, paragraph, container, text: str, style: TextStyle):
        if not text:
            return paragraph

        normalized_text = self._normalize_inline_text(text)
        if not normalized_text:
            return paragraph

        for part in self._split_text_with_paragraphs(normalized_text):
            if part == "\n":
                paragraph = self._start_new_paragraph_after(paragraph, container)
                continue
            if not part:
                continue
            if part in self._math_tokens:
                self._append_math(paragraph=paragraph, token=part, inline=not self._math_tokens[part]["display"])
                continue

            run = paragraph.add_run(part)
            self._apply_run_style(run, style)
        return paragraph

    def _append_script_text(self, paragraph, container, text: str, style: TextStyle, is_subscript: bool, is_superscript: bool):
        normalized_text = self._normalize_inline_text(text)
        if not normalized_text:
            return paragraph

        for part in self._split_text_with_paragraphs(normalized_text):
            if part == "\n":
                paragraph = self._start_new_paragraph_after(paragraph, container)
                continue
            if not part:
                continue
            run = paragraph.add_run(part)
            self._apply_run_style(run, style)
            run.font.subscript = is_subscript
            run.font.superscript = is_superscript
        return paragraph

    def _normalize_inline_text(self, text: str) -> str:
        return text.replace("\r\n", "\n").replace("\r", "\n")

    def _split_text_with_paragraphs(self, text: str) -> List[str]:
        parts: List[str] = []
        for line_index, line in enumerate(text.split("\n")):
            if line_index > 0:
                parts.append("\n")
            parts.extend(TOKEN_PATTERN.split(line))
        return parts

    def _start_new_paragraph_after(self, paragraph, container):
        new_paragraph = container.add_paragraph()
        self._copy_paragraph_format(paragraph, new_paragraph)
        return new_paragraph

    def _copy_paragraph_format(self, source, target) -> None:
        source_ppr = source._p.pPr
        if source_ppr is None:
            return

        target_ppr = target._p.pPr
        if target_ppr is not None:
            target._p.remove(target_ppr)
        target._p.insert(0, deepcopy(source_ppr))

    def _apply_run_style(self, run, style: TextStyle) -> None:
        if style.bold:
            run.bold = True
        if style.italic:
            run.italic = True
        if style.underline:
            run.underline = True

        if style.code:
            run.font.name = CODE_FONT
            run.font.size = Pt(style.size_pt or CODE_SIZE_PT)
            self._set_run_fonts(run, CODE_FONT, CODE_FONT)
            self._set_run_shading(run, "EDEDED")
        else:
            run.font.name = LATIN_FONT
            run.font.size = Pt(style.size_pt or BODY_SIZE_PT)
            self._set_run_fonts(run, LATIN_FONT, CJK_FONT)

    def _set_run_fonts(self, run, latin_font: str, east_asia_font: str) -> None:
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:ascii"), latin_font)
        r_fonts.set(qn("w:hAnsi"), latin_font)
        r_fonts.set(qn("w:eastAsia"), east_asia_font)
        r_fonts.set(qn("w:cs"), latin_font)

    def _set_run_shading(self, run, fill: str) -> None:
        r_pr = run._element.get_or_add_rPr()
        existing = r_pr.find(qn("w:shd"))
        if existing is not None:
            r_pr.remove(existing)

        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        r_pr.append(shd)

    def _set_paragraph_shading(self, paragraph, fill: str) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        existing = p_pr.find(qn("w:shd"))
        if existing is not None:
            p_pr.remove(existing)

        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        p_pr.append(shd)

    def _append_math(self, paragraph, token: str, inline: bool) -> None:
        payload = self._math_tokens.get(token)
        if payload is None:
            return

        latex = str(payload.get("latex", "")).strip()
        display = bool(payload.get("display", False))
        equation_tag: Optional[str] = None
        if display:
            latex, equation_tag = self._extract_equation_tag(latex)

        if inline:
            detached_scripts = self._extract_detached_script_tokens(latex)
            if detached_scripts:
                for marker, text in detached_scripts:
                    run = paragraph.add_run(text)
                    self._apply_run_style(run, TextStyle(size_pt=BODY_SIZE_PT))
                    run.font.superscript = marker == "^"
                    run.font.subscript = marker == "_"
                return

        try:
            omml = self.latex_to_omml(latex, display=display)
            local_tag = self._tag_name(omml)

            if inline:
                run = paragraph.add_run()
                self._set_run_fonts(run, LATIN_FONT, CJK_FONT)
                if local_tag == "omathpara":
                    math_nodes = omml.xpath(".//*[local-name()='oMath']")
                    if math_nodes:
                        run._r.append(deepcopy(math_nodes[0]))
                    else:
                        run.text = f"${latex}$"
                elif local_tag == "omath":
                        run._r.append(omml)
                else:
                    run.text = f"${latex}$"
            else:
                if equation_tag:
                    self._apply_tagged_display_math_paragraph_format(paragraph)
                    paragraph.add_run("\t")
                    self._append_omml_as_inline_run(paragraph, omml, latex)
                    run = paragraph.add_run(f"\t{equation_tag}")
                    self._apply_run_style(run, TextStyle(size_pt=BODY_SIZE_PT))
                    return
                else:
                    self._apply_display_math_paragraph_format(paragraph)
                if local_tag == "omathpara":
                    self._set_omathpara_center(omml)
                    paragraph._p.append(omml)
                elif local_tag == "omath":
                    wrapper = OxmlElement("m:oMathPara")
                    wrapper.append(omml)
                    self._set_omathpara_center(wrapper)
                    paragraph._p.append(wrapper)
                else:
                    run = paragraph.add_run()
                    self._set_run_fonts(run, LATIN_FONT, CJK_FONT)
                    run.text = f"$$ {latex} $$"
        except Exception as exc:  # noqa: BLE001
            self.last_warnings.append(f"Formula fallback: {latex[:60]} ({exc})")
            if inline:
                run = paragraph.add_run(f"${latex}$")
                self._apply_run_style(run, TextStyle(size_pt=BODY_SIZE_PT))
            else:
                self._apply_display_math_paragraph_format(paragraph)
                fallback = f"$$ {latex} $$"
                if equation_tag:
                    fallback = f"{fallback} {equation_tag}"
                run = paragraph.add_run(fallback)
                self._apply_run_style(run, TextStyle(size_pt=BODY_SIZE_PT))

    def _extract_equation_tag(self, latex: str) -> Tuple[str, Optional[str]]:
        match = re.search(r"\\tag\*?\{([^{}]+)\}\s*$", latex, flags=re.DOTALL)
        if match is None:
            return latex, None

        tag_text = match.group(1).strip()
        cleaned_latex = latex[: match.start()].rstrip()
        if not tag_text:
            return cleaned_latex, None
        if tag_text.startswith("(") and tag_text.endswith(")"):
            return cleaned_latex, tag_text
        return cleaned_latex, f"({tag_text})"

    def _append_omml_as_inline_run(self, paragraph, omml: etree._Element, latex: str) -> None:
        run = paragraph.add_run()
        self._set_run_fonts(run, LATIN_FONT, CJK_FONT)

        local_tag = self._tag_name(omml)
        if local_tag == "omathpara":
            math_nodes = omml.xpath(".//*[local-name()='oMath']")
            if math_nodes:
                run._r.append(deepcopy(math_nodes[0]))
                return
        elif local_tag == "omath":
            run._r.append(deepcopy(omml))
            return

        run.text = f"${latex}$"

    def _extract_detached_script_tokens(self, latex: str) -> Optional[List[Tuple[str, str]]]:
        stripped = latex.strip()
        if not stripped:
            return None

        pattern = re.compile(r"([_^])(?:\{([^{}]+)\}|([^\s_^{}]+))")
        tokens: List[Tuple[str, str]] = []
        position = 0

        for match in pattern.finditer(stripped):
            if match.start() != position:
                return None
            value = match.group(2) or match.group(3) or ""
            if not value:
                return None
            tokens.append((match.group(1), value))
            position = match.end()

        if position != len(stripped):
            return None
        return tokens or None

    def _apply_display_math_paragraph_format(self, paragraph) -> None:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)

    def _apply_tagged_display_math_paragraph_format(self, paragraph) -> None:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Pt(0)

        width_inches = self._page_content_width_inches()
        left_indent = paragraph.paragraph_format.left_indent
        right_indent = paragraph.paragraph_format.right_indent
        if left_indent is not None:
            width_inches -= max(0.0, left_indent.pt / 72.0)
        if right_indent is not None:
            width_inches -= max(0.0, right_indent.pt / 72.0)
        width_inches = max(1.0, width_inches)

        tab_stops = paragraph.paragraph_format.tab_stops
        try:
            tab_stops.clear_all()
        except AttributeError:
            pass
        tab_stops.add_tab_stop(Inches(width_inches / 2.0), alignment=WD_TAB_ALIGNMENT.CENTER)
        tab_stops.add_tab_stop(Inches(width_inches), alignment=WD_TAB_ALIGNMENT.RIGHT)

    def _apply_blockquote_paragraph_format(self, paragraph) -> None:
        paragraph.paragraph_format.left_indent = Pt(BLOCKQUOTE_LEFT_INDENT_PT)
        paragraph.paragraph_format.right_indent = Pt(BLOCKQUOTE_RIGHT_INDENT_PT)
        if self._paragraph_contains_display_math(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_paragraph_math_justification(paragraph, "center")
        else:
            if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        paragraph.paragraph_format.first_line_indent = Pt(BLOCKQUOTE_FIRST_LINE_INDENT_PT)
        paragraph.paragraph_format.space_before = Pt(BLOCKQUOTE_SPACE_BEFORE_PT)
        paragraph.paragraph_format.space_after = Pt(BLOCKQUOTE_SPACE_AFTER_PT)
        self._set_paragraph_shading(paragraph, BLOCKQUOTE_SHADING)
        self._set_paragraph_left_border(paragraph, BLOCKQUOTE_BORDER_COLOR)

    def _paragraph_contains_display_math(self, paragraph) -> bool:
        return bool(paragraph._p.xpath(".//*[local-name()='oMathPara']"))

    def _set_paragraph_math_justification(self, paragraph, justification: str) -> None:
        for omathpara in paragraph._p.xpath(".//*[local-name()='oMathPara']"):
            self._set_omathpara_justification(omathpara, justification)

    def _set_paragraph_left_border(self, paragraph, color: str) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)

        existing = p_bdr.find(qn("w:left"))
        if existing is not None:
            p_bdr.remove(existing)

        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), BLOCKQUOTE_BORDER_SIZE)
        left.set(qn("w:space"), str(BLOCKQUOTE_BORDER_SPACE_PT))
        left.set(qn("w:color"), color)
        p_bdr.append(left)

    def _render_horizontal_rule(self, container) -> None:
        paragraph = container.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = Pt(1)

        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)

        bottom = p_bdr.find(qn("w:bottom"))
        if bottom is not None:
            p_bdr.remove(bottom)

        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "A6A6A6")
        p_bdr.append(bottom)

    def _set_omathpara_center(self, omathpara_node: etree._Element) -> None:
        self._set_omathpara_justification(omathpara_node, "center")

    def _set_omathpara_justification(self, omathpara_node: etree._Element, justification: str) -> None:
        if self._tag_name(omathpara_node) != "omathpara":
            return

        para_pr = None
        for child in list(omathpara_node):
            if self._tag_name(child) == "omathparapr":
                para_pr = child
                break

        if para_pr is None:
            para_pr = OxmlElement("m:oMathParaPr")
            omathpara_node.insert(0, para_pr)

        jc = None
        for child in list(para_pr):
            if self._tag_name(child) == "jc":
                jc = child
                break
        if jc is None:
            jc = OxmlElement("m:jc")
            para_pr.append(jc)
        jc.set(qn("m:val"), justification)

    def _insert_image(self, paragraph, image_node: etree._Element, config: AppConfig) -> None:
        src = image_node.get("src", "").strip()
        temp_paths: List[Path] = []

        try:
            if self._is_remote_url(src):
                resolved = self._download_remote_image(src)
                temp_paths.append(resolved)
            else:
                resolved = self.resolve_image_path(src, config.asset_root)

            if resolved is None or not resolved.exists():
                if src:
                    warning = f"图片未找到: {src}（资源根目录: {config.asset_root}）"
                    self.last_warnings.append(warning)
                    run = paragraph.add_run(f"[{warning}]")
                    self._apply_run_style(run, TextStyle(size_pt=BODY_SIZE_PT))
                return

            picture_path = resolved
            if resolved.suffix.lower() == ".svg":
                picture_path, temp_path = self._convert_svg_to_png(resolved)
                temp_paths.append(temp_path)

            width_inches = self._parse_image_width(image_node.get("width"), picture_path)
            run = paragraph.add_run()
            if width_inches:
                run.add_picture(str(picture_path), width=Inches(width_inches))
            else:
                run.add_picture(str(picture_path))
        except Exception as exc:  # noqa: BLE001
            detail = self._error_detail(exc)
            suffix = Path(urllib.parse.urlparse(src).path).suffix.lower() if self._is_remote_url(src) else Path(src).suffix.lower()
            if suffix == ".svg":
                detail = f"SVG 图片自动转换失败，请安装 cairosvg/Cairo 或先转为 PNG/JPG。{detail}"
            warning = f"图片插入失败: {src}（{detail}）"
            self.last_warnings.append(warning)
            run = paragraph.add_run(f"[{warning}]")
            self._apply_run_style(run, TextStyle(size_pt=BODY_SIZE_PT))
        finally:
            for temp_path in temp_paths:
                try:
                    temp_path.unlink(missing_ok=True)
                except Exception:
                    pass

    def _is_remote_url(self, src: str) -> bool:
        return src.lower().startswith(("http://", "https://"))

    def _download_remote_image(self, src: str) -> Path:
        request = urllib.request.Request(src, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(request, timeout=20) as response:
            data = response.read()
            content_type = (response.headers.get("Content-Type") or "").split(";", 1)[0].strip().lower()

        suffix = Path(urllib.parse.urlparse(src).path).suffix.lower()
        if not suffix:
            suffix = {
                "image/jpeg": ".jpg",
                "image/jpg": ".jpg",
                "image/png": ".png",
                "image/gif": ".gif",
                "image/bmp": ".bmp",
                "image/tiff": ".tiff",
                "image/svg+xml": ".svg",
            }.get(content_type, ".img")

        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
            temp_file.write(data)
            return Path(temp_file.name)

    def _convert_svg_to_png(self, image_path: Path) -> Tuple[Path, Path]:
        try:
            import cairosvg
        except Exception as exc:  # noqa: BLE001
            detail = self._error_detail(exc)
            raise RuntimeError(f"cairosvg 或 Cairo 图形库不可用，无法自动转换 SVG: {detail}") from exc

        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_file:
            temp_path = Path(temp_file.name)

        try:
            cairosvg.svg2png(url=str(image_path), write_to=str(temp_path))
        except Exception:
            temp_path.unlink(missing_ok=True)
            raise

        return temp_path, temp_path

    def _error_detail(self, exc: Exception, max_length: int = 260) -> str:
        detail = re.sub(r"\s+", " ", str(exc).strip()) or exc.__class__.__name__
        if len(detail) > max_length:
            return f"{detail[:max_length]}..."
        return detail

    def _parse_image_width(self, width_attr: Optional[str], image_path: Path) -> Optional[float]:
        max_width = self._page_content_width_inches()
        if not width_attr:
            return self._natural_or_max_width(image_path, max_width)

        raw = width_attr.strip().lower()
        try:
            if raw.endswith("%"):
                return max(0.2, min(max_width * max(min(float(raw.rstrip("%")), 100.0), 1.0) / 100.0, max_width))
            if raw.endswith("px"):
                return max(0.2, min(float(raw[:-2]) / 96.0, max_width))
            if raw.endswith("cm"):
                return max(0.2, min(float(raw[:-2]) / 2.54, max_width))
            if raw.endswith("mm"):
                return max(0.2, min(float(raw[:-2]) / 25.4, max_width))
            if raw.endswith("in"):
                return max(0.2, min(float(raw[:-2]), max_width))
            if raw.endswith("pt"):
                return max(0.2, min(float(raw[:-2]) / 72.0, max_width))
            return max(0.2, min(float(raw) / 96.0, max_width))
        except ValueError:
            return self._natural_or_max_width(image_path, max_width)

    def _natural_or_max_width(self, image_path: Path, max_width: float) -> float:
        try:
            with Image.open(image_path) as img:
                dpi = img.info.get("dpi", (96, 96))[0] or 96
                natural_width = img.width / float(dpi)
                return max(0.2, min(natural_width, max_width))
        except Exception:
            return max_width

    def _page_content_width_inches(self) -> float:
        if self._document is None:
            return 6.0
        section = self._document.sections[0]
        return float(section.page_width - section.left_margin - section.right_margin) / 914400.0

    def _extract_list_start(self, node: etree._Element) -> int:
        raw_start = (node.get("start") or "1").strip()
        try:
            return max(1, int(raw_start))
        except ValueError:
            return 1

    def _configure_list_paragraph(self, paragraph, level: int, marker_text: str) -> None:
        marker_width_pt = self._estimate_list_marker_width(marker_text)
        left_indent = Pt((LIST_INDENT_PT * max(level - 1, 0)) + marker_width_pt)
        paragraph.paragraph_format.left_indent = left_indent
        paragraph.paragraph_format.first_line_indent = Pt(-marker_width_pt)
        paragraph.paragraph_format.tab_stops.add_tab_stop(left_indent)

    def _bullet_prefix(self, level: int) -> str:
        bullets = ["•", "◦", "▪"]
        return bullets[min(max(level - 1, 0), len(bullets) - 1)]

    def _estimate_list_marker_width(self, marker_text: str) -> float:
        return max(LIST_MARKER_MIN_WIDTH_PT, len(marker_text) * LIST_MARKER_CHAR_WIDTH_PT)

    def _add_paragraph_with_style(self, container, style_names: Iterable[str]):
        for style_name in style_names:
            try:
                return container.add_paragraph(style=style_name)
            except KeyError:
                continue
        return container.add_paragraph()

    def _apply_alignment(self, paragraph, node: etree._Element) -> None:
        align = self._extract_alignment(node)
        if align is not None:
            paragraph.alignment = align

    def _extract_alignment(self, node: etree._Element) -> Optional[WD_ALIGN_PARAGRAPH]:
        align_attr = (node.get("align") or "").strip().lower()
        style_attr = (node.get("style") or "").lower()

        if "text-align" in style_attr:
            match = re.search(r"text-align\s*:\s*(left|center|right|justify)", style_attr)
            if match:
                align_attr = match.group(1)

        mapping = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        return mapping.get(align_attr)

    def _build_output_filename(self, markdown_text: str, config: AppConfig) -> str:
        cleaned = re.sub(r"```[\s\S]*?```|~~~[\s\S]*?~~~", " ", markdown_text)
        cleaned = re.sub(r"\$\$.+?\$\$", " ", cleaned, flags=re.DOTALL)
        cleaned = re.sub(r"\\\[.+?\\\]", " ", cleaned, flags=re.DOTALL)
        cleaned = re.sub(r"\\\(.+?\\\)", " ", cleaned)
        cleaned = re.sub(r"`[^`\n]+`", " ", cleaned)
        cleaned = re.sub(r"<[^>]+>", " ", cleaned)
        cleaned = re.sub(r"[\r\n\t]+", " ", cleaned)
        cleaned = re.sub(r"\s+", " ", cleaned).strip()

        normalized = re.sub(r"[^\w\u4e00-\u9fff]+", "", cleaned, flags=re.UNICODE)
        base = normalized[: max(1, config.title_chars)] or "document"

        if config.auto_timestamp:
            return f"{base}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        return f"{base}.docx"

    def _single_block_math_token(self, node: etree._Element) -> Optional[str]:
        if len(node) != 0:
            return None

        text = (node.text or "").strip()
        if text in self._math_tokens and bool(self._math_tokens[text]["display"]):
            return text
        return None

    def _is_caption_text(self, text: str) -> bool:
        normalized = re.sub(r"\s+", " ", text or "").strip().lower()
        return bool(CAPTION_PREFIX_PATTERN.match(normalized))

    def _is_abstract_title(self, text: str) -> bool:
        normalized = re.sub(r"\s+", "", text or "")
        return normalized in {"摘要", "摘要:", "摘要："}

    def _is_image_only_block(self, node: etree._Element) -> bool:
        if not node.xpath(".//*[local-name()='img']"):
            return False

        text = node.text_content() or ""
        return not text.strip()

    def _tag_name(self, node) -> str:
        if node is None:
            return ""
        if isinstance(node, str):
            return node
        if not hasattr(node, "tag"):
            return ""
        if isinstance(node.tag, str):
            return node.tag.split("}")[-1].lower()
        return ""

    def _markdown_heading_style_name(self, level: int) -> str:
        if level <= 1:
            return ARTICLE_TITLE_STYLE
        return f"Heading {min(level - 1, 5)}"

    def _markdown_heading_size(self, level: int) -> float:
        return {
            1: ARTICLE_TITLE_SIZE_PT,
            2: H1_SIZE_PT,
            3: H2_SIZE_PT,
            4: H3_SIZE_PT,
            5: H4_SIZE_PT,
            6: H5_SIZE_PT,
        }.get(level, BODY_SIZE_PT)

    def _clear_cell(self, cell) -> None:
        tc = cell._tc
        for paragraph in cell.paragraphs:
            tc.remove(paragraph._p)


def build_converter() -> MarkdownToDocxConverter:
    return MarkdownToDocxConverter()
